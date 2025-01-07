import streamlit as st
from azure_utils import read_pdf
from neo4j_utils import query_neo4j
from embeddings import generate_embedding, openai
import io
from docx import Document

async def minute_meeting_checker_section(llm, driver):
    st.title("Minute Meeting Checker")
    st.subheader("Upload Credentials and Meeting Minutes for Analysis")

    # Input Microsoft Document Intelligence credentials
    with st.expander("Enter Microsoft Document Intelligence Credentials"):
        endpoint = st.text_input("Endpoint:", placeholder="Enter your Microsoft Document Intelligence endpoint")
        key = st.text_input("API Key:", placeholder="Enter your Microsoft Document Intelligence API key", type="password")

    if not endpoint or not key:
        st.warning("Please provide your Microsoft Document Intelligence credentials to proceed.")
        return

    # File upload for meeting minutes
    with st.expander("Upload Meeting Minutes File"):
        minutes_file = st.file_uploader(
            "Upload a minutes file (PDF, DOCX, TXT, or CSV):",
            type=["pdf", "docx", "txt", "csv"],
            key="minutes"
        )

    if not minutes_file:
        st.info("Please upload a meeting minutes file to begin analysis.")
        return

    try:
        # Step 1: Extract Text from Uploaded File
        st.subheader("Step 1: Extract Text from Uploaded File")
        minutes_content = read_pdf(minutes_file, endpoint, key)
        if minutes_content.strip():
            st.text_area("Extracted Text", value=minutes_content, height=300)
        else:
            st.error("The uploaded file does not contain readable text.")
            return

        if len(minutes_content.split()) < 10:
            st.error("The content appears too short or invalid for analysis.")
            return

        # Step 2: Analyze the Meeting Minutes
        st.subheader("Step 2: Analyze the Meeting Minutes")
        analysis_prompt = f"""
        คุณเป็นผู้ช่วย AI สำหรับสรุปรายงานการประชุมสำหรับผู้ตรวจสอบของธนาคารแห่งประเทศไทยเป็นภาษาไทย
        กรุณาจัดทำรายงานที่มีโครงสร้างชัดเจนและครบถ้วนโดยแบ่งเป็นหัวข้อหลักดังนี้:

        **วาระการประชุม:**
        - [หัวข้อย่อย]: [รายละเอียด เช่น เงื่อนไขที่กำหนด,อัตราร้อยละ, จำนวนเงิน, ระยะเวลา และกลุ่มเป้าหมายที่ชัดเจน]

        **ประเด็นสำคัญอื่นๆ ที่ได้มีการหารือ:**
        - [หัวข้อย่อย]: [รายละเอียด เช่น เงื่อนไขที่กำหนด,อัตราร้อยละ, จำนวนเงิน, ระยะเวลา และกลุ่มเป้าหมายที่ชัดเจน ]

        **งานที่ต้องดำเนินการและผู้รับผิดชอบ:**
        - [หัวข้อย่อย]: [รายละเอียด เช่น ขั้นตอนปฏิบัติ และชื่อบุคคล/หน่วยงานที่รับผิดชอบ]

        ข้อมูลที่ควรระบุในทุกหัวข้อ:
        1. รายละเอียดตัวเลข เช่น ร้อยละ, จำนวนเงิน, หรือข้อมูลที่เจาะจงเกี่ยวกับมาตรการ
        2. ระยะเวลา เช่น วันเริ่มต้น, วันสิ้นสุด หรือระยะเวลาการบังคับใช้
        3. กลุ่มเป้าหมายหรือเงื่อนไข เช่น ประเภทลูกค้า, ข้อจำกัด, หรือผลกระทบ
        4. ถ้ามีชื่อเต็มของตัวย่อให้เขียนแบบนี้เสมอ ชื่อเต็ม(ตัวย่อ)
        5. ถ้ามีหัวข้อแยกย่อยในหัวข้อย่อยมากกว่า1ให้ใช้เป็นข้อ 1.,2.,3. ... ถ้ามีข้อย่อยลงไปอีกให้ใช้ 1.1,1.2,1.3, ...
        6. ห้ามมีเครื่องหมาย"-"นำหน้าถ้าไม่ใช่ - [หัวข้อย่อย]:

        รายงานการประชุม:
        {minutes_content}
        
        ตอบในรูปแบบที่กำหนดเพื่อให้สามารถแยกและประมวลผลหัวข้อย่อยได้ง่าย โดยใช้โครงสร้าง:
        - [หัวข้อย่อย]: [รายละเอียด]
        """
        
        try:
            analysis_response = llm._call(analysis_prompt).strip()
            if not analysis_response:
                st.error("The LLM did not return any analysis. Please check the input or LLM configuration.")
                return
            st.text_area("Analyzed Content", value=analysis_response, height=300)
        except Exception as e:
            st.error(f"An error occurred during analysis: {e}")
            return

        # Step 3: Split and Process Subtopics
        st.subheader("Step 3: Split and Process Subtopics")
        sections = {"วาระการประชุม": [], "ประเด็นสำคัญอื่นๆ ที่ได้มีการหารือ": [], "งานที่ต้องดำเนินการและผู้รับผิดชอบ": []}

        try:
            for section in sections.keys():
                start_index = analysis_response.find(section)
                if start_index != -1:
                    next_sections = [analysis_response.find(next_section, start_index + len(section))
                                     for next_section in sections.keys() if next_section != section]
                    next_sections = [idx for idx in next_sections if idx != -1]
                    end_index = min(next_sections) if next_sections else len(analysis_response)

                    section_content = analysis_response[start_index + len(section):end_index].strip()

                    subtopics = []
                    for line in section_content.split("\n"):
                        line = line.strip()
                        if line.startswith("- "):
                            subtopic, content = line.split(":", 1) if ":" in line else (line, "")
                            subtopic = subtopic.replace("- ", "").strip()
                            subtopics.append({"subtopic": subtopic, "content": content.strip()})
                        elif subtopics and line:
                            subtopics[-1]["content"] += f" {line}"

                    sections[section] = subtopics
                else:
                    st.warning(f"Section '{section}' was not found in the analysis response.")

            # Step 4: Match Subtopics to Documents
            st.subheader("Step 4: Match Subtopics to Documents")

            # Query to retrieve document names
            document_query = """
            MATCH (d:Document)
            RETURN d.name AS documentName
            """
            document_results = query_neo4j(driver, document_query, {})

            # Normalize and trim document names
            import unicodedata
            document_names = [
                unicodedata.normalize('NFC', result['documentName'].strip())
                for result in document_results
            ]

            # Log retrieved document names for debugging
            #st.write("Retrieved Document Names:", document_names)

            # Initialize dictionary to store matched documents
            matched_documents = {}

            for section, subtopics in sections.items():
                for subtopic in subtopics:
                    subtopic_name = subtopic.get("subtopic", "Unknown Subtopic")
                    subtopic_content = subtopic.get("content", "")

                    # Updated matching prompt to ensure only document names are returned
                    match_prompt = f"""
                    คุณเป็นผู้ช่วย AI สำหรับตรวจสอบบันทึกการประชุม
                    หัวข้อย่อย: {subtopic_name}
                    เนื้อหา: {subtopic_content}
                    รายการชื่อเอกสารที่มีอยู่:
                    {', '.join(document_names)}

                    โปรดจับคู่หัวข้อย่อยกับชื่อเอกสารที่เหมาะสมที่สุดจากรายการด้านบน และตอบในรูปแบบ:
                    <ชื่อเอกสาร>

                    ตัวอย่าง:
                    การกำหนดหลักเกณฑ์ วิธีการ และเงื่อนไขในการประกอบธุรกิจบัตรเครดิตของธนาคารพานิชย์
                    """
                    
                    try:
                        matched_document = llm._call(match_prompt).strip()
                    except Exception as e:
                        st.error(f"An error occurred during LLM matching: {e}")
                        matched_document = "No Match"

                    # Log the matched document for debugging
                    #st.write(f"Matched Document for '{subtopic_name}': {matched_document}")

                    # Store the matched document name directly
                    matched_documents[subtopic_name] = matched_document

                # Step 5: Process Subtopics with Matched Documents
                import uuid
                # Define report_data BEFORE the button
                report_data = []

            if st.button("Continue", key="unique_run_step_5"):
                st.subheader("Step 5: Process Subtopics with Matched Documents")
                report_data.clear()

                # Enumerate the sections to use 'i' in the key
                for i, (section, subtopics) in enumerate(sections.items()):
                    if section in ["งานที่ต้องดำเนินการและผู้รับผิดชอบ", "ประเด็นสำคัญอื่นๆ ที่ได้มีการหารือ"]:
                        #st.write(f"Skipping section: {section}")
                        continue  # Skip further processing for these sections

                    with st.expander(section):
                        # Enumerate the subtopics to use 'j' in the key
                        for j, subtopic in enumerate(subtopics):
                            subtopic_name = subtopic.get("subtopic", "Unknown Subtopic")
                            subtopic_content = subtopic.get("content", "")
                            matched_document = matched_documents.get(subtopic_name, "No Match")

                            # Clean and extract document names
                            document_names = [
                                doc.strip() for doc in matched_document.split(",")
                                if "เอกสารที่เกี่ยวข้อง" not in doc and doc.strip()
                            ]

                            st.write(f"**{subtopic_name}**: {subtopic_content} "
                                    f"\n(Matched Document(s): {', '.join(document_names)})")

                            # Enumerate the document_names to use 'k' in the key
                            for k, document_name in enumerate(document_names):
                                if not subtopic_content:
                                    continue

                                # Generate embedding, etc.
                                topic_embedding = await generate_embedding(subtopic_content)
                                if isinstance(topic_embedding, str) and topic_embedding.startswith("Error"):
                                    st.error(f"Error generating embedding for {subtopic_name}: {topic_embedding}")
                                    continue

                                cypher_query = """
                                MATCH (chunk:Chunk)-[:PART_OF]->(doc:Document)
                                WHERE trim(doc.name) = trim($matchedDocument)
                                WITH doc, genai.vector.encode($userInput, "OpenAI", { token: $apiKey }) AS userEmbedding, chunk
                                CALL db.index.vector.queryNodes('chunk_embedding_index', 5, userEmbedding)
                                YIELD node AS matchedChunk, score
                                WITH DISTINCT matchedChunk, score, doc  
                                OPTIONAL MATCH (matchedChunk)-[:NEXT]->(nextChunk:Chunk)
                                OPTIONAL MATCH (previousChunk:Chunk)-[:NEXT]->(matchedChunk)
                                RETURN 
                                    matchedChunk.text AS chunkText, 
                                    nextChunk.text AS nextChunkText, 
                                    previousChunk.text AS previousChunkText,
                                    doc.name AS documentName, 
                                    doc.related_sections AS documentRelatedSections, 
                                    score
                                ORDER BY score DESC
                                """
                            parameters = {
                                "userInput": f"{subtopic_name}: {subtopic_content}",
                                "apiKey": openai.api_key,
                                "matchedDocument": document_name.strip()
                            }

                            # Create truly unique Streamlit keys by combining i, j, k (indices) + optional UUID:
                            unique_key_query = f"cypher_{i}_{j}_{k}_{uuid.uuid4()}"
                            unique_key_results = f"results_{i}_{j}_{k}_{uuid.uuid4()}"

                            # Mask the API key for display
                            masked_api_key = openai.api_key[:4] + "****" + openai.api_key[-4:]
                            debug_cypher_query = cypher_query.replace("$apiKey", masked_api_key)

                            st.text_area(
                                f"Cypher Query for {subtopic_name} ({document_name})",
                                value=debug_cypher_query,
                                height=150,
                                key=unique_key_query
                            )

                            # Execute the query
                            results = query_neo4j(driver, cypher_query, parameters)

                            # Optional: show query results
                            st.text_area(
                                 f"Query Results for {subtopic_name} ({document_name})",
                                 value=str(results),
                                 height=200,
                                 key=unique_key_results
                             )
                                
                            comment_prompt = f"""
                            คุณเป็นผู้ช่วย AI สำหรับตรวจสอบบันทึกการประชุม
                            Subtopic: {subtopic_name}
                            Content: {subtopic_content}

                            Based on the retrieved database results:
                            {results}

                            โปรดให้ความเห็นตามแนวทางใดแนวทางหนึ่งจากแนวทางต่อไปนี้:
                                1. หากไม่พบข้อขัดแย้ง ให้ระบุว่า "ไม่พบข้อขัดแย้งต่อประกาศธนาคาร" และไม่ต้องให้ความเห็นเพิ่มเติม
                                2. หากข้อมูลไม่เพียงพอ ให้ระบุว่า "ข้อมูลไม่เพียงพอ" และแจ้งข้อมูลที่ควรขอเพิ่มเติมจากธนาคารเจ้าของ minute meeting และประเมินว่าข้อมูลretrieved database resultsเกี่ยวข้องกับsubtopicนี้หรือไม่
                                3. หากพบข้อขัดแย้ง ให้ระบุว่า "พบข้อขัดแย้ง" พร้อมมาตรากฎหมายที่เกี่ยวข้องและชื่อประกาศที่เกี่ยวข้อง ถ้ามีบางส่วนของหัวข้อย่อยที่มีข้อมูลไม่เพียงพอให้เขียนกำกับที่ส่วนนั้นว่าข้อมูลไม่เพียงพอ และแจ้งข้อมูลที่ควรขอเพิ่มเติมจากธนาคารเจ้าของ minute meeting
                                4. หากข้อมูลจาก Retrieved database results ไม่เกี่ยวข้องกับ หัวข้อย่อย ให้ระบุว่า "Retrieved database results ไม่เกี่ยวข้องกับหัวข้อย่อย"
                            ให้ระบุมาตรากฏหมายที่เกี่ยวข้องกับประกาศธนาคารฉบับนั้นในทุกวาระ (ซึ่งจะระบุอยู่หลัง "documentRelatedSections :" ในรูปแบบของList จากresults)

                            ตัวอย่างโครงสร้าง:
                            **วาระที่ 2: การพิจารณามาตรการส่งเสริมการใช้จ่ายผ่านบัตรเครดิต:**
                            
                            มาตรการที่ 1: การลดเกณฑ์การชำระหนี้ขั้นต่ำเหลือ 5% ของยอดคงค้าง (จากเดิม 10%)
                            ข้อขัดแย้ง: ประกาศธนาคารแห่งประเทศไทยที่ สนส. 90/2563 เรื่อง การกำหนดหลักเกณฑ์ วิธีการ และเงื่อนไขในการประกอบธุรกิจบัตรเครดิต (30 กรกฎาคม 2563) กำหนดให้ผู้ประกอบธุรกิจบัตรเครดิตต้องเรียกเก็บชำระหนี้ขั้นต่ำไม่น้อยกว่า 10% ของยอดคงค้างในแต่ละเดือน การลดเกณฑ์เหลือ 5% จึงขัดแย้งกับประกาศฉบับนี้
                            มาตรากฎหมายที่เกี่ยวข้อง: พระราชบัญญัติธุรกิจสถาบันการเงิน พ.ศ. 2551 มาตรา 41, 71 (ระบุเป็นตัวเลขมาตราที่เกี่ยวข้องทั้งหมด)
                            ชื่อประกาศที่เกี่ยวข้อง: ประกาศธนาคารแห่งประเทศไทยที่ สนส. 90/2563 เรื่อง การกำหนดหลักเกณฑ์ วิธีการ และเงื่อนไขในการประกอบธุรกิจบัตรเครดิต (ไม่ต้องกล่าวถึงประกาศและหนังสือเวียนที่ยกเลิก)

                            มาตราที่ 2: ....
                            """
                            llm_comment = llm._call(comment_prompt).strip()

                            # Add processed subtopic data
                            report_data.append((subtopic_name, subtopic_content, results, document_name, llm_comment))

        except Exception as e:
            st.error(f"An error occurred while splitting and processing subtopics: {e}")
            return

        # Step 6: Generate Compliance Report
        st.subheader("Step 6: Generate Compliance Report")

        try:
            # 1. Initialize once
            violations_report = "### Violations Report\n"
            retrieved_report = "### Retrieved Database Results\n"

            # 2. Loop through all subtopics
            for subtopic_name, subtopic_content, results, matched_document, llm_comment in report_data:
                # Append to violations_report (Markdown formatting)
                violations_report += f"\n#### {subtopic_name}\n"
                violations_report += f"Matched Document: {matched_document}\n"
                violations_report += f"Comment: {llm_comment}\n"
                
                # (Optional) Similarly add to retrieved_report if needed
                # if results:
                #     retrieved_report += f"\n####**Subtopic**: {subtopic_name}\n"
                #     retrieved_report += f"  **Content**: {subtopic_content}\n"
                #     retrieved_report += f"  **Matched Document**: {matched_document}\n"
                #     for result in results:
                #         retrieved_report += f"  - **Chunk**: {result['chunkText']}\n"
                #         retrieved_report += f"    **Score**: {result['score']}\n"

            # Display them in Streamlit
            st.markdown(violations_report, unsafe_allow_html=True)
            # st.text_area("Retrieved Results", value=retrieved_report, height=300)

            # 3. Build a docx file from the same violations_report text
            doc = Document()

            # 4. Add a heading for Analyzed Content
            doc.add_heading("Analyzed Content", level=2)
            # 5. Add the actual analysis response
            doc.add_paragraph(analysis_response)

            # 6. Add a heading for the Violations Report
            doc.add_heading("Violations Report", level=2)
            # 7. Add the violations report
            doc.add_paragraph(violations_report)

            # 8. Convert the docx document to an in-memory buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # 9. Provide the download button for the docx file
            st.download_button(
                label="Download Violations Report (docx)",
                data=buffer,
                file_name="violations_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error(f"An error occurred while generating the compliance report: {e}")


    except Exception as e:
        st.error(f"An error occurred during processing: {e}")
